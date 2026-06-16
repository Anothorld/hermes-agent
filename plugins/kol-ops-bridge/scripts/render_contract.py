"""Render a POVISON influencer contract docx from a JSON field dict.

Usage (CLI):
    python3 render_contract.py \
        --template /path/to/povison_agreement.docx \
        --output   /path/to/out.docx \
        --fields   /path/to/fields.json   # or "-" for stdin

The renderer:
  1. Replaces every ``${KEY}`` placeholder in body paragraphs and table cells
     using a flat key map derived from the field dict.
  2. Removes the "A one-time flat fee of ${FEE} USD." paragraph when
     ``fields["fee"]`` is null/absent.
  3. Fills the deliverable row in the first table (the row that follows the
     "should create original content..." sentence). When more than one
     deliverable is given, additional rows are cloned from the original.

Placeholders are split across docx runs by Word's formatting; we merge runs
within each placeholder-containing paragraph onto run 0 (preserving its style)
before substitution.
"""

from __future__ import annotations

import argparse
import copy
import datetime as _dt
import json
import logging
import re
import sys
import time
from pathlib import Path
from typing import Any, Iterable, Mapping

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Row

log = logging.getLogger("kol_ops_bridge.render_contract")

PLACEHOLDER_RE = re.compile(r"\$\{[A-Z_][A-Z0-9_ ]*\}")
FEE_PLACEHOLDER = "${FEE}"
# Template placeholders are styled red in povison_agreement.docx (EE0000).
_PLACEHOLDER_COLORS = frozenset({"EE0000", "FF0000", "C00000", "RED"})
_DEBUG_LOG = Path("/Users/arnold/agent_prj/.cursor/debug-88a275.log")


def _debug_log(hypothesis_id: str, location: str, message: str, data: dict[str, Any]) -> None:
    # region agent log
    try:
        payload = {
            "sessionId": "88a275",
            "hypothesisId": hypothesis_id,
            "location": location,
            "message": message,
            "data": data,
            "timestamp": int(time.time() * 1000),
        }
        with _DEBUG_LOG.open("a", encoding="utf-8") as fh:
            fh.write(json.dumps(payload, ensure_ascii=False) + "\n")
    except OSError:
        pass
    # endregion


def _run_color_val(run) -> str | None:
    r_pr = run._element.find(qn("w:rPr"))
    if r_pr is None:
        return None
    color = r_pr.find(qn("w:color"))
    if color is None:
        return None
    return (color.get(qn("w:val")) or "").upper()


def _clear_placeholder_highlight(run) -> None:
    """Drop template placeholder red after a value is substituted."""
    r_pr = run._element.find(qn("w:rPr"))
    if r_pr is None:
        return
    color = r_pr.find(qn("w:color"))
    if color is None:
        return
    val = (color.get(qn("w:val")) or "").upper()
    if val in _PLACEHOLDER_COLORS:
        r_pr.remove(color)


def _flatten_fields(fields: Mapping[str, Any]) -> dict[str, str]:
    influencer = fields.get("influencer") or {}
    product = fields.get("product") or {}
    fee = fields.get("fee") or {}

    fee_str = ""
    if fee:
        amount = fee.get("amount")
        currency = (fee.get("currency") or "USD").strip()
        if amount not in (None, "", 0):
            fee_str = f"{amount}" if not currency or currency == "USD" else f"{amount} {currency}"

    return {
        "DATE": str(fields.get("date") or _dt.date.today().isoformat()),
        "INFLUENCER_FULL NAME": str(influencer.get("full_name") or ""),
        "INFLUENCER_EMAIL": str(influencer.get("email") or ""),
        "INFLUENCER_PHONE NUMBER": str(influencer.get("phone") or ""),
        "INFLUENCER_ADDRESS": str(influencer.get("address") or ""),
        "INSTAGRAM_LINK": str(influencer.get("instagram") or ""),
        "TIKTOK_LINK": str(influencer.get("tiktok") or ""),
        "YOUTUBE_LINK": str(influencer.get("youtube") or ""),
        "PRODUCT_SPECS": str(product.get("specs") or ""),
        "PRODUCT_LINK": str(product.get("link") or ""),
        "FEE": fee_str,
    }


def _merge_runs_to_first(paragraph) -> None:
    """Collapse all run text onto run[0]; blank the others."""
    runs = paragraph.runs
    if len(runs) <= 1:
        return
    merged = "".join(r.text for r in runs)
    runs[0].text = merged
    for r in runs[1:]:
        r.text = ""


def _substitute_in_paragraph(paragraph, mapping: Mapping[str, str]) -> None:
    if "${" not in paragraph.text:
        return
    _merge_runs_to_first(paragraph)
    if not paragraph.runs:
        return
    run = paragraph.runs[0]
    before_color = _run_color_val(run)
    text = run.text
    substituted = False
    for key, val in mapping.items():
        needle = "${" + key + "}"
        if needle in text:
            text = text.replace(needle, val)
            substituted = True
    # Wipe any stray ${...} that lacked a mapping so the contract doesn't leak placeholders.
    text = PLACEHOLDER_RE.sub("", text)
    run.text = text
    if substituted and before_color in _PLACEHOLDER_COLORS:
        _clear_placeholder_highlight(run)
        _debug_log(
            "H1",
            "render_contract.py:_substitute_in_paragraph",
            "cleared placeholder red after substitution",
            {
                "before_color": before_color,
                "after_color": _run_color_val(run),
                "text_preview": text[:80],
            },
        )


def _iter_all_paragraphs(doc) -> Iterable[Any]:
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _remove_fee_line(doc) -> None:
    for paragraph in list(doc.paragraphs):
        if FEE_PLACEHOLDER in paragraph.text:
            element = paragraph._element
            element.getparent().remove(element)
            return


def _fill_deliverable_row(row: _Row, deliverable: Mapping[str, Any]) -> None:
    columns = (
        str(deliverable.get("type") or ""),
        str(deliverable.get("description") or ""),
        str(deliverable.get("quantity") or ""),
        str(deliverable.get("requirements") or deliverable.get("length") or ""),
        str(deliverable.get("time_of_uploading") or ""),
        str(deliverable.get("platform_of_uploading") or ""),
    )
    for cell, value in zip(row.cells, columns):
        # one cell can have several paragraphs; we keep only the first
        # and clear the rest so we don't smuggle leftover template text.
        first = cell.paragraphs[0]
        for extra in cell.paragraphs[1:]:
            extra._element.getparent().remove(extra._element)
        _merge_runs_to_first(first)
        if first.runs:
            run = first.runs[0]
            before_color = _run_color_val(run)
            run.text = value
            if before_color in _PLACEHOLDER_COLORS:
                _clear_placeholder_highlight(run)
        else:
            first.add_run(value)


def _populate_deliverables(doc, deliverables: list[Mapping[str, Any]]) -> None:
    """Replace every data row (row 1+) with one row per ``deliverables`` entry.

    All non-header rows in the original template are treated as sample data and
    overwritten — the table after rendering contains exactly the rows the
    campaign asked for.
    """
    if not deliverables or not doc.tables:
        return
    table = doc.tables[0]
    if len(table.rows) < 2:
        return

    template_tr = copy.deepcopy(table.rows[1]._tr)
    # Drop every data row (everything after the header).
    for row in list(table.rows[1:]):
        row._tr.getparent().remove(row._tr)

    header_tr = table.rows[0]._tr
    insert_after = header_tr
    for _ in deliverables:
        clone = copy.deepcopy(template_tr)
        insert_after.addnext(clone)
        insert_after = clone

    for row, deliverable in zip(table.rows[1:], deliverables):
        _fill_deliverable_row(row, deliverable)


def render(template_path: Path, output_path: Path, fields: Mapping[str, Any]) -> Path:
    mapping = _flatten_fields(fields)
    _debug_log(
        "H1",
        "render_contract.py:render",
        "render start",
        {
            "product_specs_preview": str((fields.get("product") or {}).get("specs") or "")[:80],
            "template": str(template_path),
        },
    )
    doc = Document(str(template_path))

    if not fields.get("fee"):
        _remove_fee_line(doc)

    for paragraph in _iter_all_paragraphs(doc):
        _substitute_in_paragraph(paragraph, mapping)

    deliverables = fields.get("deliverables") or []
    if isinstance(deliverables, list) and deliverables:
        _populate_deliverables(doc, deliverables)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    red_runs = []
    for paragraph in _iter_all_paragraphs(doc):
        for run in paragraph.runs:
            color = _run_color_val(run)
            if color in _PLACEHOLDER_COLORS and (run.text or "").strip():
                red_runs.append({"color": color, "text": (run.text or "")[:80]})
    _debug_log(
        "H1",
        "render_contract.py:render",
        "render complete",
        {"output": str(output_path), "remaining_red_runs": red_runs[:10], "red_run_count": len(red_runs)},
    )
    return output_path


def _load_fields(source: str) -> dict[str, Any]:
    if source == "-":
        return json.load(sys.stdin)
    return json.loads(Path(source).read_text(encoding="utf-8"))


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--fields", required=True, help='JSON file path, or "-" for stdin')
    args = parser.parse_args(argv)

    fields = _load_fields(args.fields)
    out = render(args.template, args.output, fields)
    print(str(out))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
